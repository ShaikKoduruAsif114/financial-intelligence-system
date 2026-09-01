import csv
import pickle
import re
from pathlib import Path
from typing import Any, Dict, Iterable, List

import chromadb
import fitz
import pdfplumber
from chromadb.utils import embedding_functions
from docx import Document as DocxDocument
from dotenv import load_dotenv
from langchain_text_splitters import RecursiveCharacterTextSplitter
from rank_bm25 import BM25Okapi

from vision_analysis import summarize_chart_image

load_dotenv()

DATA_DIR = Path("data/raw")
CHUNK_SIZE = 600
CHUNK_OVERLAP = 80
COLLECTION_NAME = "financial_docs"
SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".doc", ".xlsx", ".xls", ".csv", ".txt", ".png", ".jpg", ".jpeg"}


def normalize_whitespace(value: str) -> str:
    return re.sub(r"\s+", " ", value or "").strip()


def table_to_markdown(rows: Iterable[Iterable[str]]) -> str:
    cleaned = []
    for row in rows:
        if not row:
            continue
        cleaned.append([str(cell).strip() if cell is not None else "" for cell in row])
    if not cleaned:
        return ""

    headers = cleaned[0]
    rows_md = ["| " + " | ".join(headers) + " |", "| " + " | ".join(["---"] * len(headers)) + " |"]
    for row in cleaned[1:]:
        pad = row + [""] * max(0, len(headers) - len(row))
        rows_md.append("| " + " | ".join(pad[: len(headers)]) + " |")
    return "\n".join(rows_md)


def extract_pdf_chunks(pdf_path: Path) -> List[Dict[str, Any]]:
    """Convert a PDF into text, table, and figure chunks while retaining page/section metadata."""
    chunks: List[Dict[str, Any]] = []

    with pdfplumber.open(str(pdf_path)) as pdf:
        for page_number, page in enumerate(pdf.pages, start=1):
            cleaned_text = normalize_whitespace(page.extract_text() or "")
            if cleaned_text:
                chunks.append(
                    {
                        "kind": "text",
                        "content": cleaned_text,
                        "source": pdf_path.name,
                        "page": page_number,
                        "section": "Body text",
                    }
                )

            tables = page.extract_tables() or []
            for table_index, table in enumerate(tables, start=1):
                table_md = table_to_markdown(table)
                if table_md:
                    chunks.append(
                        {
                            "kind": "table",
                            "content": f"Table {table_index} (Page {page_number})\n{table_md}",
                            "source": pdf_path.name,
                            "page": page_number,
                            "section": "Tables",
                        }
                    )

            fig_captions = re.findall(r"(Figure\s+\d+.*?(?:\.|\n|$))", page.extract_text() or "", flags=re.IGNORECASE)
            for figure_index, caption in enumerate(fig_captions, start=1):
                clean_caption = normalize_whitespace(caption)
                if clean_caption:
                    chunks.append(
                        {
                            "kind": "figure",
                            "content": clean_caption,
                            "source": pdf_path.name,
                            "page": page_number,
                            "section": "Figures",
                        }
                    )

    # Add native PDF page text blocks from PyMuPDF as a second pass when needed for structure-rich docs.
    with fitz.open(pdf_path) as doc:
        for page_number in range(len(doc)):
            page = doc[page_number]
            blocks = page.get_text("blocks")
            for block in blocks:
                block_text = normalize_whitespace(block[4])
                if len(block_text) < 30:
                    continue
                if not any(existing["source"] == pdf_path.name and existing["page"] == page_number + 1 and existing["kind"] == "text" and existing["content"] == block_text for existing in chunks):
                    chunks.append(
                        {
                            "kind": "text",
                            "content": block_text,
                            "source": pdf_path.name,
                            "page": page_number + 1,
                            "section": "Body text",
                        }
                    )
    return chunks


def docx_to_markdown(doc_path: Path) -> List[Dict[str, Any]]:
    doc = DocxDocument(str(doc_path))
    chunks: List[Dict[str, Any]] = []

    paragraphs = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    if paragraphs:
        chunks.append({"kind": "text", "content": paragraphs, "source": doc_path.name, "page": 1, "section": "Document text"})

    for table_index, table in enumerate(doc.tables, start=1):
        rows = [[cell.text for cell in row.cells] for row in table.rows]
        md = table_to_markdown(rows)
        if md:
            chunks.append({"kind": "table", "content": f"Table {table_index}\n{md}", "source": doc_path.name, "page": 1, "section": "Tables"})
    return chunks


def spreadsheet_to_markdown(file_path: Path) -> List[Dict[str, Any]]:
    try:
        import pandas as pd
    except ImportError as exc:  # pragma: no cover
        raise RuntimeError("pandas is required to ingest spreadsheet files.") from exc

    sheets = pd.read_excel(file_path, sheet_name=None)
    chunks: List[Dict[str, Any]] = []
    for sheet_name, dataframe in sheets.items():
        table_md = dataframe.to_markdown(index=False)
        if table_md:
            chunks.append({"kind": "table", "content": f"Sheet: {sheet_name}\n{table_md}", "source": file_path.name, "page": 1, "section": "Tables"})
    return chunks


def csv_to_markdown(file_path: Path) -> List[Dict[str, Any]]:
    rows = []
    with file_path.open("r", encoding="utf-8-sig", newline="") as csvfile:
        reader = csv.reader(csvfile)
        for row in reader:
            if any(cell.strip() for cell in row):
                rows.append(row)
    markdown = table_to_markdown(rows)
    return [{"kind": "table", "content": markdown, "source": file_path.name, "page": 1, "section": "Tables"}] if markdown else []


def text_to_chunks(file_path: Path) -> List[Dict[str, Any]]:
    content = file_path.read_text(encoding="utf-8", errors="ignore")
    return [{"kind": "text", "content": normalize_whitespace(content), "source": file_path.name, "page": 1, "section": "Document text"}] if content.strip() else []


def chart_image_to_chunks(file_path: Path) -> List[Dict[str, Any]]:
    summary = summarize_chart_image(file_path)
    if not summary:
        return []
    return [{
        "kind": "figure",
        "content": f"Chart image: {file_path.name}\n{summary}",
        "source": file_path.name,
        "page": 1,
        "section": "Figure analysis",
    }]


def parse_document(file_path: Path) -> List[Dict[str, Any]]:
    suffix = file_path.suffix.lower()
    if suffix == ".pdf":
        return extract_pdf_chunks(file_path)
    if suffix in {".docx", ".doc"}:
        return docx_to_markdown(file_path)
    if suffix in {".xlsx", ".xls"}:
        return spreadsheet_to_markdown(file_path)
    if suffix == ".csv":
        return csv_to_markdown(file_path)
    if suffix == ".txt":
        return text_to_chunks(file_path)
    if suffix in {".png", ".jpg", ".jpeg"}:
        return chart_image_to_chunks(file_path)
    return []


def build_document_chunks(sources: List[Dict[str, Any]], chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP) -> List[Dict[str, Any]]:
    """Split extracted document elements into semantically meaningful chunks with type metadata."""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", "|", ".", " "],
        length_function=len,
    )

    chunks: List[Dict[str, Any]] = []
    for index, source in enumerate(sources):
        content = str(source.get("content", "")).strip()
        if not content:
            continue

        splits = splitter.split_text(content)
        for split_index, split in enumerate(splits):
            chunk_id = source.get("chunk_id") or f"{source.get('source', 'doc')}_{source.get('page', 1)}_{source.get('kind', 'text')}_{index}_{split_index}"
            chunks.append(
                {
                    "id": chunk_id,
                    "content": normalize_whitespace(split),
                    "kind": source.get("kind", "text"),
                    "source": source.get("source", "unknown"),
                    "page": source.get("page", 1),
                    "section": source.get("section", "Document text"),
                    "chunk_index": split_index,
                }
            )
    return chunks


def collect_documents(data_dir: Path) -> List[Path]:
    if not data_dir.exists():
        return []
    return sorted(path for path in data_dir.rglob("*") if path.is_file() and path.suffix.lower() in SUPPORTED_EXTENSIONS)


def store_in_chromadb(chunks: List[Dict[str, Any]], vectorstore_path: str = "vectorstore"):
    client = chromadb.PersistentClient(path=vectorstore_path)
    ef = embedding_functions.SentenceTransformerEmbeddingFunction(model_name="BAAI/bge-small-en-v1.5")
    collection = client.get_or_create_collection(name=COLLECTION_NAME, embedding_function=ef)

    batch_size = 100
    for i in range(0, len(chunks), batch_size):
        batch = chunks[i : i + batch_size]
        collection.add(
            documents=[chunk["content"] for chunk in batch],
            ids=[chunk["id"] for chunk in batch],
            metadatas=[
                {
                    "source": chunk["source"],
                    "page": chunk["page"],
                    "kind": chunk["kind"],
                    "section": chunk["section"],
                }
                for chunk in batch
            ],
        )

    tokenized_corpus = [chunk["content"].lower().split(" ") for chunk in chunks]
    bm25 = BM25Okapi(tokenized_corpus)
    with open(f"{vectorstore_path}/bm25_index.pkl", "wb") as handle:
        pickle.dump({"bm25": bm25, "chunks": chunks}, handle)

    return collection


def ingest_documents(data_dir: Path = DATA_DIR, chunk_size: int = CHUNK_SIZE, chunk_overlap: int = CHUNK_OVERLAP):
    documents = collect_documents(data_dir)
    if not documents:
        print(f"No supported document files found in {data_dir}. Add PDFs, DOCX, XLSX, CSV, or TXT files to ingest them.")
        return []

    all_sources: List[Dict[str, Any]] = []
    for doc_path in documents:
        print(f"Processing: {doc_path.name}")
        parsed = parse_document(doc_path)
        all_sources.extend(parsed)

    chunks = build_document_chunks(all_sources, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    if not chunks:
        print("No chunks were generated from the provided files.")
        return []

    print(f"Generated {len(chunks)} structured chunks from {len(documents)} documents.")
    store_in_chromadb(chunks)
    return chunks


def main():
    ingest_documents()


if __name__ == "__main__":
    main()
